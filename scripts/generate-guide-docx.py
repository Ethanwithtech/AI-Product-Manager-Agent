#!/usr/bin/env python3
"""Generate the PM Team Hub configuration and usage guide as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level, size, color in [(1, 18, "165DFF"), (2, 14, "333333"), (3, 12, "555555")]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor.from_string(color)
    hs.font.name = "Arial"
    hs.paragraph_format.space_before = Pt(12)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "165DFF",
            qn("w:val"): "clear",
        })
        shading.append(bg)
    # Rows
    for row_data in rows:
        r = t.add_row()
        for i, val in enumerate(row_data):
            r.cells[i].text = val
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    shading = run._element.get_or_add_rPr()
    return p


def add_note(text, prefix="💡 提示："):
    p = doc.add_paragraph()
    run = p.add_run(prefix)
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    return p


# ═══════════════════════════════════════════════
# Title Page
# ═══════════════════════════════════════════════
doc.add_paragraph()  # spacing
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("PM Team Hub")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor.from_string("165DFF")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("配置与使用指南")
r.font.size = Pt(20)
r.font.color.rgb = RGBColor.from_string("333333")

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run("更新日期：2026年3月25日 · 版本 v1.2")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor.from_string("999999")

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. 系统架构
# ═══════════════════════════════════════════════
doc.add_heading("1. 系统架构概览", level=1)
doc.add_paragraph(
    "PM Team Hub 采用「Skills + MCP + 轻量管理台」三层分布式架构，专为 5 人以下 PM 团队设计。"
)
add_table(
    ["层级", "组件", "说明"],
    [
        ["① AI Skills", "63 个 Skill + 6 个 Agent", "每个 PM 在自己的 CodeBuddy 中加载，AI 帮你写需求、生成 UI、同步进度、分析反馈"],
        ["② MCP Server", "FastAPI + SQLite + ChromaDB", "共享数据中心，所有 PM 连接同一个服务器，12 个 MCP 工具"],
        ["③ 管理台", "React + Arco Design", "浏览器中管理知识库、需求单、进度看板、反馈分析、规则模板"],
    ],
    col_widths=[3.5, 5.5, 8.5],
)

doc.add_heading("双数据库架构", level=2)
add_table(
    ["数据库", "用途", "存储内容"],
    [
        ["SQLite", "结构化数据", "需求单、进度、模板、反馈、知识库元数据"],
        ["ChromaDB", "向量数据库", "知识库文档的语义嵌入（all-MiniLM-L6-v2），支持自然语言检索"],
    ],
    col_widths=[3, 4, 10.5],
)

# ═══════════════════════════════════════════════
# 2. 环境准备
# ═══════════════════════════════════════════════
doc.add_heading("2. 环境准备与启动", level=1)

doc.add_heading("2.1 后端服务", level=2)
doc.add_paragraph("① 进入后端目录：cd mcp-server")
doc.add_paragraph("② 安装依赖：pip install -r requirements.txt")
doc.add_paragraph("   新增依赖：pypdf（PDF 解析）、python-docx（Word 解析）")
doc.add_paragraph("③ 启动服务：python main.py")
doc.add_paragraph("④ 团队模式：./scripts/start-server.sh --team")
add_note("团队模式会自动打印服务器 IP 和 MCP 配置模板，方便其他 PM 复制使用。")

doc.add_heading("2.2 前端管理台", level=2)
doc.add_paragraph("① 进入前端目录：cd platform/web")
doc.add_paragraph("② 安装依赖：npm install")
doc.add_paragraph("③ 本地开发：npm run dev")
doc.add_paragraph("④ 连接远程后端：VITE_API_URL=http://<服务器IP>:8000 npm run dev")

doc.add_heading("2.3 团队共享部署", level=2)
doc.add_paragraph(
    "核心原则：所有 PM 连接同一个 MCP Server，共享知识库、需求库、进度板。"
    "找一台内网服务器（或用你的 Mac 当服务器），部署后端，其他人通过内网 IP 访问。"
)

# ═══════════════════════════════════════════════
# 3. CodeBuddy 配置
# ═══════════════════════════════════════════════
doc.add_heading("3. CodeBuddy 配置", level=1)

doc.add_heading("3.1 MCP 连接", level=2)
doc.add_paragraph("在项目根目录创建 .mcp.json 文件：")
add_code_block(
    '{\n'
    '  "mcpServers": {\n'
    '    "pm-team-hub": {\n'
    '      "type": "sse",\n'
    '      "url": "http://<服务器IP>:8000/mcp/sse"\n'
    '    }\n'
    '  }\n'
    '}'
)
add_note("将 <服务器IP> 替换为实际地址。团队所有人使用同一个 IP。")

doc.add_heading("3.2 Skills 加载", level=2)
doc.add_paragraph(
    "63 个 Skills 已通过符号链接配置在 .codebuddy/skills/ 目录下。"
    "CodeBuddy 打开项目时自动识别加载，无需额外操作。"
)

doc.add_heading("3.3 TAPD 认证配置（可选）", level=2)
doc.add_paragraph("如需使用需求单同步到 TAPD 功能，创建凭证文件 ~/.tapd/credentials：")
add_code_block("access_token=<你的TAPD访问令牌>\nenv=OA")
add_note("env 可选 OA（内网）或 IDC（外网）。", "⚠️ 注意：")

doc.add_page_break()

# ═══════════════════════════════════════════════
# 4. 功能模块
# ═══════════════════════════════════════════════
doc.add_heading("4. 功能模块详解", level=1)

# 4.1 知识库
doc.add_heading("4.1 产品知识库 📚", level=2)
doc.add_paragraph(
    "内置 ChromaDB 向量数据库，上传文档自动分块（1000字/块，200字重叠）、嵌入、索引。"
    "AI 写需求时自动调用 search_knowledge 检索相关内容。"
)
add_table(
    ["项目", "说明"],
    [
        ["支持格式", ".md、.txt、.pdf（新增）、.docx（新增）、.html（新增）— 拖拽上传自动解析"],
        ["推荐上传", "IM/TCCC 产品文档、竞品分析、用户调研报告、技术方案、接口文档"],
        ["MCP 工具", "search_knowledge、add_knowledge_document、list_knowledge_documents、delete_knowledge_document"],
        ["检索测试", "点击「检索测试」按钮，输入自然语言查询，相关度分数越接近 1 越好"],
        ["Embedding", "all-MiniLM-L6-v2 模型，完全本地化，不依赖外部 API"],
    ],
    col_widths=[3.5, 14],
)

# 4.2 需求单
doc.add_heading("4.2 需求单管理 + TAPD 同步 📝", level=2)
doc.add_paragraph(
    "AI 生成 PRD 后存入本地 MCP Server，可选同步到 TAPD。完整 4 阶段流程："
)
add_table(
    ["阶段", "内容", "MCP 工具 / 脚本"],
    [
        ["Phase 1: 上下文组装", "检索知识库 + 读取模板 + 查看进度板", "search_knowledge, get_templates, get_progress_board"],
        ["Phase 2: AI 生成 PRD", "基于上下文 + 模板生成结构化需求文档", "—（AI 内部处理）"],
        ["Phase 3: 存储 & 冲突检测", "存入本地 + 自动检测全队冲突", "create_requirement, check_conflicts"],
        ["Phase 4: TAPD 同步（新增）", "创建 TAPD 需求 → 上传附件 → 嵌入图片", "stories_create + upload-attachment.py + upload-image.py"],
    ],
    col_widths=[4, 6, 7.5],
)
add_note("Phase 4 需要先配置 ~/.tapd/credentials。同步前 AI 会询问确认。", "⚠️ ")

# 4.3 进度看板
doc.add_heading("4.3 进度看板 📊", level=2)
doc.add_paragraph("4 列看板：规划中 → 进行中 → 评审中 → 已完成。")
doc.add_paragraph(
    "自动冲突检测：检测两种冲突类型 —— ① 模块重叠（两人修改同一模块）"
    " ② 关键词重叠（≥3 个相同关键词）。创建需求时自动触发。"
)
add_note(
    "所有 PM 连接同一个 MCP Server 时，创建需求会自动扫描全队正在进行的工作，"
    "实时通知冲突。"
)

# 4.4 反馈分析
doc.add_heading("4.4 反馈分析 💬", level=2)
doc.add_paragraph(
    "收集标书反馈、App Store 评论、用户访谈、调查问卷等多渠道数据。"
    "AI 自动提取主题、情感评分、影响等级评估，生成产品迭代建议。"
)
doc.add_paragraph("在 CodeBuddy 中使用 feedback-insight-engine Skill，可与 PM 互动讨论迭代优先级。")

# 4.5 模板管理
doc.add_heading("4.5 规则与模板 📋", level=2)
doc.add_paragraph("三种分类：需求模板、团队规则、检查清单。")
p = doc.add_paragraph()
r = p.add_run("两种创建方式（新增文件上传）：")
r.bold = True
doc.add_paragraph("① 点击「新建模板」手动编写 Markdown 内容")
doc.add_paragraph("② 点击「上传模板文件」导入 .md/.txt/.docx/.pdf 文件，内容自动解析为模板")
add_note("AI 生成需求时会自动调用 get_templates() 读取你的模板，按模板格式生成 PRD。")

doc.add_page_break()

# ═══════════════════════════════════════════════
# 5. Skills 索引
# ═══════════════════════════════════════════════
doc.add_heading("5. PM 团队专属 Skills（17个）", level=1)
doc.add_paragraph("以下是为 PM 团队协作场景专门开发的 Skills：")

add_table(
    ["Skill 名称", "类型", "功能说明"],
    [
        ["product-knowledge-base", "Component", "组织和查询产品文档，为 AI 提供上下文"],
        ["requirement-generator", "Workflow", "从自然语言生成 PRD，可同步 TAPD（Phase 4）"],
        ["ui-draft-generator", "Workflow", "基于需求生成 React + Arco Design UI 代码"],
        ["product-sync-agent", "Interactive", "同步产品进度，检测团队工作冲突"],
        ["fullchain-efficiency", "Workflow", "知识检索 → PRD → UI 稿 → 团队同步，全自动"],
        ["feedback-insight-engine", "Interactive", "收集分析标书/评论/访谈，生成迭代建议"],
        ["skill-navigator", "Component", "路由索引，列出全部 63 个 Skill 并按场景分类"],
        ["tapd-toolkit", "Component", "TAPD 图片/附件上传、查询与下载"],
        ["multi-perspective-evaluation", "Interactive", "五维度全面评估方案"],
        ["asr-sentence-recognition", "Component", "腾讯云 ASR 语音识别：短音频/极速/异步"],
        ["docx", "Component", "创建、读取、编辑 Word 文档"],
        ["pdf", "Component", "PDF 读取、合并、拆分、OCR、加密"],
        ["tencent-meeting-mcp", "Interactive", "腾讯会议管理：预约/修改/取消、录制转写"],
        ["wecom-doc-skills-v2", "Component", "企业微信文档读取：文档/表格/幻灯片/思维导图"],
        ["wecom-message", "Component", "通过 Webhook 向企业微信发送消息"],
        ["woa-preview", "Workflow", "将文档/报告发布到 pages.woa.com 预览"],
        ["xiaohongshu-mcp", "Interactive", "小红书内容自动化：发布、搜索、分析笔记"],
    ],
    col_widths=[4.5, 2.5, 10.5],
)

# ═══════════════════════════════════════════════
# 6. 推荐工作流
# ═══════════════════════════════════════════════
doc.add_heading("6. 推荐日常工作流", level=1)

add_table(
    ["时机", "操作", "使用的 Skill / 命令"],
    [
        ["🌅 早上", "同步进度，检查冲突", "product-sync-agent"],
        ["📝 写需求时", "一键生成 PRD + 同步 TAPD", "/generate-requirement → TAPD Phase 4"],
        ["💬 收到反馈时", "AI 分析反馈，推荐迭代方向", "feedback-insight-engine"],
        ["🎨 出 UI 稿时", "基于需求生成 React 原型", "/generate-ui-draft"],
        ["⚡ 全链路", "需求到 UI 一步到位", "/fullchain"],
        ["📚 上传文档", "知识库中上传产品文档", "管理台「知识库」页面 或 add_knowledge_document MCP 工具"],
        ["📋 管理模板", "创建/上传需求模板", "管理台「规则与模板」页面"],
    ],
    col_widths=[3, 5, 9.5],
)

# ═══════════════════════════════════════════════
# 7. MCP 工具一览
# ═══════════════════════════════════════════════
doc.add_heading("7. MCP 工具一览（12 个）", level=1)
add_table(
    ["类别", "工具名", "说明"],
    [
        ["知识库", "search_knowledge", "语义搜索知识库文档"],
        ["知识库", "add_knowledge_document", "添加文档到知识库"],
        ["知识库", "list_knowledge_documents", "列出所有知识库文档"],
        ["知识库", "delete_knowledge_document", "删除知识库文档"],
        ["需求单", "create_requirement", "创建需求（自动触发冲突检测）"],
        ["需求单", "get_requirement", "获取单个需求详情"],
        ["需求单", "list_requirements", "列出所有需求"],
        ["需求单", "update_requirement", "更新需求内容/状态"],
        ["进度", "update_progress", "更新工作进度"],
        ["进度", "get_progress_board", "获取进度看板"],
        ["进度", "check_conflicts", "检测工作冲突"],
        ["模板", "get_templates", "获取模板（AI 自动调用）"],
    ],
    col_widths=[2.5, 5, 10],
)

# ═══════════════════════════════════════════════
# 8. 常见问题
# ═══════════════════════════════════════════════
doc.add_heading("8. 常见问题 FAQ", level=1)

doc.add_heading("Q: 其他 PM 怎么连接到 MCP Server？", level=3)
doc.add_paragraph(
    "A: 把 MCP Server 部署到一台团队共享的内网服务器上，"
    "所有人的 .mcp.json 指向该服务器 IP。使用 ./scripts/start-server.sh --team 可自动打印配置。"
)

doc.add_heading("Q: 创建需求后能自动通知其他人冲突吗？", level=3)
doc.add_paragraph(
    "A: 是的。调用 create_requirement() 时系统自动执行 check_conflicts()，"
    "扫描所有 planning/in_progress/review 状态的需求，检测模块重叠和关键词重叠，实时返回冲突列表。"
)

doc.add_heading("Q: 知识库支持哪些文件格式？", level=3)
doc.add_paragraph(
    "A: 支持 .md、.txt、.pdf、.docx、.html 五种格式。"
    "上传后自动解析为纯文本，分块嵌入 ChromaDB 向量数据库。"
)

doc.add_heading("Q: TAPD 同步需要什么前提条件？", level=3)
doc.add_paragraph(
    "A: 需要在 ~/.tapd/credentials 文件中配置 access_token 和 env。"
    "生成 PRD 后 AI 会询问是否同步，不会自动推送。"
)

doc.add_heading("Q: 模板怎么导入？", level=3)
doc.add_paragraph(
    "A: 两种方式：① 在「规则与模板」页面点击「新建模板」手动编写 "
    "② 点击「上传模板文件」直接导入 .md/.txt/.docx/.pdf 文件。"
)

# ── Save ──
output_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "docs",
    "PM团队CodeBuddy配置与使用指南.docx",
)
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"✅ Word 文档已生成: {output_path}")
