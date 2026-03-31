"""File parser service — extract text from PDF, Word, Markdown, and plain text files.

Supports: .pdf, .docx, .md, .txt, .html
Used by knowledge base upload and template file upload.
"""

from __future__ import annotations

import io
import re
from typing import Optional


def parse_file(content: bytes, filename: str) -> str:
    """Parse file content based on file extension. Returns plain text."""
    ext = _get_extension(filename).lower()

    if ext in (".md", ".txt", ".text", ".markdown"):
        return content.decode("utf-8", errors="replace")

    if ext == ".pdf":
        return _parse_pdf(content)

    if ext in (".docx", ".doc"):
        return _parse_docx(content)

    if ext in (".html", ".htm"):
        return _parse_html(content)

    # Fallback: try UTF-8 decode
    return content.decode("utf-8", errors="replace")


def get_doc_type(filename: str) -> str:
    """Infer doc_type from filename extension."""
    ext = _get_extension(filename).lower()
    type_map = {
        ".md": "md",
        ".markdown": "md",
        ".txt": "txt",
        ".text": "txt",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".html": "html",
        ".htm": "html",
    }
    return type_map.get(ext, "txt")


def _get_extension(filename: str) -> str:
    """Extract file extension from filename."""
    if "." in filename:
        return "." + filename.rsplit(".", 1)[-1]
    return ""


def _parse_pdf(content: bytes) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n\n".join(pages)
    except ImportError:
        raise RuntimeError(
            "pypdf is not installed. Install with: pip install pypdf>=5.0.0"
        )
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")


def _parse_docx(content: bytes) -> str:
    """Extract text from Word document using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure with markdown-style headers
                if para.style and para.style.name.startswith("Heading"):
                    level = 1
                    try:
                        level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    paragraphs.append(f"{'#' * level} {text}")
                else:
                    paragraphs.append(text)

        # Also extract table content
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))

        return "\n\n".join(paragraphs)
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Install with: pip install python-docx>=1.0.0"
        )
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")


def _parse_html(content: bytes) -> str:
    """Extract text from HTML by stripping tags (simple approach)."""
    text = content.decode("utf-8", errors="replace")
    # Remove script and style blocks
    text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    # Replace common block tags with newlines
    text = re.sub(r"<(?:br|p|div|h[1-6]|li|tr)[^>]*>", "\n", text, flags=re.IGNORECASE)
    # Strip remaining tags
    text = re.sub(r"<[^>]+>", "", text)
    # Clean up whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# Supported extensions for validation
SUPPORTED_EXTENSIONS = {".md", ".markdown", ".txt", ".text", ".pdf", ".docx", ".doc", ".html", ".htm"}


def is_supported(filename: str) -> bool:
    """Check if a file type is supported for parsing."""
    return _get_extension(filename).lower() in SUPPORTED_EXTENSIONS
